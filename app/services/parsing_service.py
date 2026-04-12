"""Document parsing — Lane 1: PyMuPDF, Lane 2: Unstructured fallback."""
import logging
import io
from dataclasses import dataclass, field
from typing import List, Optional

logger = logging.getLogger(__name__)


@dataclass
class ParsedPage:
    page_number: int
    text: str
    section_title: Optional[str] = None


@dataclass
class ParsedDocument:
    pages: List[ParsedPage] = field(default_factory=list)
    language: str = "en"
    parse_method: str = "unknown"

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text.strip())


def _parse_with_pymupdf(file_bytes: bytes, file_type: str) -> Optional[ParsedDocument]:
    """Lane 1 — fast, for standard digital PDFs and TXT."""
    try:
        import fitz  # PyMuPDF

        if file_type == "txt":
            text = file_bytes.decode("utf-8", errors="replace")
            return ParsedDocument(
                pages=[ParsedPage(page_number=1, text=text)],
                parse_method="pymupdf_txt",
            )

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                pages.append(ParsedPage(page_number=i + 1, text=text))
        doc.close()

        if not pages:
            return None  # trigger fallback

        return ParsedDocument(pages=pages, parse_method="pymupdf")
    except Exception as e:
        logger.warning(f"PyMuPDF parsing failed: {e}")
        return None


def _parse_with_unstructured(file_bytes: bytes, file_type: str, file_name: str) -> ParsedDocument:
    """Lane 2 — robust fallback for scanned PDFs, DOCX, etc."""
    try:
        from unstructured.partition.auto import partition
        from unstructured.documents.elements import Title, NarrativeText, ListItem, Table

        elements = partition(
            file=io.BytesIO(file_bytes),
            metadata_filename=file_name,
        )

        pages: dict[int, list[str]] = {}
        current_page = 1
        current_title = None

        for el in elements:
            page_num = getattr(el.metadata, "page_number", None) or current_page
            current_page = page_num

            if isinstance(el, Title):
                current_title = str(el)
                text = f"## {el}\n"
            else:
                text = str(el)

            if page_num not in pages:
                pages[page_num] = []
            pages[page_num].append(text)

        parsed_pages = [
            ParsedPage(page_number=pn, text="\n".join(texts))
            for pn, texts in sorted(pages.items())
            if "\n".join(texts).strip()
        ]

        return ParsedDocument(pages=parsed_pages, parse_method="unstructured")
    except Exception as e:
        logger.error(f"Unstructured parsing failed: {e}")
        raise RuntimeError(f"All parsing strategies failed for file '{file_name}': {e}")


def _parse_docx(file_bytes: bytes) -> ParsedDocument:
    """DOCX via python-docx (fast, no Unstructured needed for simple DOCX)."""
    try:
        import docx

        doc = docx.Document(io.BytesIO(file_bytes))
        pages: list[ParsedPage] = []
        current_texts: list[str] = []
        page_number = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            current_texts.append(text)

        # DOCX has no page concept — treat whole doc as page 1
        if current_texts:
            pages.append(ParsedPage(page_number=1, text="\n".join(current_texts)))

        return ParsedDocument(pages=pages, parse_method="python_docx")
    except Exception as e:
        logger.warning(f"python-docx parsing failed: {e}")
        return None


async def parse_document(file_bytes: bytes, file_type: str, file_name: str) -> ParsedDocument:
    """
    Entry point. Tries the fastest parser first and falls back to Unstructured.
    """
    import asyncio

    def _parse():
        if file_type == "txt":
            result = _parse_with_pymupdf(file_bytes, file_type)
            return result or ParsedDocument(
                pages=[ParsedPage(page_number=1, text=file_bytes.decode("utf-8", errors="replace"))],
                parse_method="raw_text",
            )

        if file_type == "docx":
            result = _parse_docx(file_bytes)
            if result and result.pages:
                return result
            # fallback
            return _parse_with_unstructured(file_bytes, file_type, file_name)

        if file_type == "pdf":
            result = _parse_with_pymupdf(file_bytes, file_type)
            if result and result.pages:
                return result
            logger.info(f"PyMuPDF returned empty output for '{file_name}', trying Unstructured.")
            return _parse_with_unstructured(file_bytes, file_type, file_name)

        raise ValueError(f"Unsupported file type: {file_type}")

    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _parse)
